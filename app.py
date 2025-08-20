from flask import Flask, request, jsonify
import requests
import openai
from docx import Document
import os

app = Flask(__name__)
openai.api_key = "Т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"   # сюда вставь свой ключ с platform.openai.com

# создаём документ, если его ещё нет
if not os.path.exists("answers.docx"):
    doc = Document()
    doc.add_heading("Опросник", 0)
    doc.save("answers.docx")

@app.route("/voice", methods=["POST"])
def handle_voice():
    data = request.json
    file_url = data.get("file_url")   # BotHelp пришлёт ссылку на голосовое
    question = data.get("question")

    # скачиваем голосовое
    voice_file = requests.get(file_url).content
    with open("voice.ogg", "wb") as f:
        f.write(voice_file)

    # транскрипция через OpenAI Whisper
    with open("voice.ogg", "rb") as audio:
        transcript = openai.audio.transcriptions.create(
            model="whisper-1",
            file=audio,
            language="ru"
        )
    text = transcript.text

    # сохраняем в Word
    doc = Document("answers.docx")
    doc.add_paragraph(f"{question}: {text}")
    doc.save("answers.docx")

    # возвращаем текст обратно в BotHelp
    return jsonify({"text": text})
